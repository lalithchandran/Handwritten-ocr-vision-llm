import streamlit as st
import requests
import io
import fitz
from docx import Document
from fpdf import FPDF
from datetime import datetime

# -----------------------------
# CONFIG
# -----------------------------
FASTAPI_URL = "http://localhost:8000/extract-ocr"

st.set_page_config(
    page_title="AI Handwritten OCR",
    page_icon="🧠",
    layout="wide"
)

# -----------------------------
# CLEAN PROFESSIONAL STYLING
# -----------------------------
st.markdown("""
<style>

/* Remove default large padding */
.block-container {
    padding-top: 2rem !important;
    padding-bottom: 1rem !important;
}

/* Background */
.main {
    background: linear-gradient(135deg, #0b1220, #0f172a);
    color: #f8fafc;
}

/* Sidebar */
section[data-testid="stSidebar"] {
    background-color: #111827 !important;
}
section[data-testid="stSidebar"] * {
    color: #e5e7eb !important;
}

/* Hero */
.hero {
    background: linear-gradient(135deg, #6366f1, #8b5cf6);
    padding: 45px 40px;
    border-radius: 18px;
    margin-bottom: 30px;
}
.hero h1 {
    font-size: 40px;
    font-weight: 800;
    margin-bottom: 12px;
    color: white;
}
.hero p {
    font-size: 18px;
    color: #e0e7ff;
}

/* Metric Cards */
.metric-card {
    background: #1e293b;
    padding: 22px;
    border-radius: 16px;
    text-align: center;
    box-shadow: 0 8px 25px rgba(0,0,0,0.45);
}
.metric-card h4 {
    color: #cbd5e1;
    font-size: 15px;
    font-weight: 600;
}
.metric-card h2 {
    font-size: 32px;
    font-weight: 800;
    margin-top: 10px;
    color: #ffffff;
}

/* Content Cards */
.card {
    background: #1f2937;
    padding: 22px;
    border-radius: 16px;
    box-shadow: 0 8px 25px rgba(0,0,0,0.4);
}

/* File uploader */
section[data-testid="stFileUploader"] {
    background-color: #1f2937;
    padding: 18px !important;
    border-radius: 14px;
    border: 1px solid #334155;
    margin-bottom: 8px !important;
}

/* Reduce gap between uploader and filename */
div[data-testid="stFileUploader"] + div {
    margin-top: -8px !important;
}

/* Buttons */
.stButton>button {
    border-radius: 12px;
    height: 3em;
    font-weight: 700;
    background: linear-gradient(135deg, #6366f1, #8b5cf6);
    color: white;
    border: none;
}

/* Download buttons */
.stDownloadButton>button {
    background-color: #334155;
    color: white;
    border-radius: 10px;
    font-weight: 600;
}

/* Textarea */
textarea {
    background-color: #0f172a !important;
    color: #f8fafc !important;
    border-radius: 12px !important;
    border: 1px solid #334155 !important;
}

/* Reduce column gap */
.css-ocqkz7 {
    gap: 1rem !important;
}

</style>
""", unsafe_allow_html=True)

# -----------------------------
# UTILITIES
# -----------------------------
def convert_pdf_to_image(pdf_bytes):
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    page = doc.load_page(0)
    pix = page.get_pixmap(dpi=150)
    return pix.tobytes("jpeg")

def create_word_file(text):
    doc = Document()
    doc.add_heading('Extracted Handwritten Text', 0)
    doc.add_paragraph(text)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def create_pdf_file(text):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("helvetica", size=12)
    safe_text = text.encode('latin-1', 'replace').decode('latin-1')
    pdf.multi_cell(0, 8, safe_text)
    return bytes(pdf.output())

# -----------------------------
# SIDEBAR
# -----------------------------
with st.sidebar:
    st.title("🧠 AI OCR")
    st.markdown("Premium Dashboard")
    st.divider()
    st.success("Backend: Connected")
    st.info("Model: LLaMA.cpp")
    st.divider()
    st.caption("© 2026 Lalith AI Labs")

# -----------------------------
# HERO
# -----------------------------
st.markdown("""
<div class="hero">
    <h1>AI Handwritten OCR Platform</h1>
    <p>Upload handwritten images or PDFs. Extract, edit, and export text instantly using AI.</p>
</div>
""", unsafe_allow_html=True)

# -----------------------------
# SESSION STATE
# -----------------------------
if "extracted_text" not in st.session_state:
    st.session_state.extracted_text = ""

if "files_processed" not in st.session_state:
    st.session_state.files_processed = 0

# -----------------------------
# METRICS
# -----------------------------
m1, m2, m3 = st.columns(3)

with m1:
    st.markdown(f"""
    <div class="metric-card">
        <h4>📂 Files Processed</h4>
        <h2>{st.session_state.files_processed}</h2>
    </div>
    """, unsafe_allow_html=True)

with m2:
    st.markdown("""
    <div class="metric-card">
        <h4>⚡ Status</h4>
        <h2>Ready</h2>
    </div>
    """, unsafe_allow_html=True)

with m3:
    st.markdown(f"""
    <div class="metric-card">
        <h4>🕒 Session</h4>
        <h2>{datetime.now().strftime("%M:%S")}</h2>
    </div>
    """, unsafe_allow_html=True)

st.markdown("<br>", unsafe_allow_html=True)

# -----------------------------
# FILE UPLOAD
# -----------------------------
uploaded_file = st.file_uploader(
    "Upload Image or PDF",
    type=["png", "jpg", "jpeg", "pdf"],
    label_visibility="collapsed"
)

# -----------------------------
# MAIN CONTENT
# -----------------------------
if uploaded_file is not None:

    col1, col2 = st.columns(2)

    # LEFT SIDE - PREVIEW
    with col1:
        st.markdown('<div class="card">', unsafe_allow_html=True)
        st.subheader("📄 Document Preview")

        file_bytes = uploaded_file.read()

        if uploaded_file.type == "application/pdf":
            image_bytes = convert_pdf_to_image(file_bytes)
            st.image(image_bytes, use_container_width=True)
        else:
            image_bytes = file_bytes
            st.image(image_bytes, use_container_width=True)

        if st.button("🚀 Extract Text", use_container_width=True):
            with st.spinner("Analyzing handwriting..."):
                try:
                    files = {"file": ("image.jpg", image_bytes, "image/jpeg")}
                    response = requests.post(FASTAPI_URL, files=files)

                    if response.status_code == 200:
                        data = response.json()
                        st.session_state.extracted_text = data.get("extracted_text", "")
                        st.session_state.files_processed += 1
                        st.success("Extraction Complete!")
                    else:
                        st.error("Backend Error")

                except requests.exceptions.ConnectionError:
                    st.error("Backend not running on port 8000.")

        st.markdown('</div>', unsafe_allow_html=True)

    # RIGHT SIDE - TEXT
    with col2:
        st.markdown('<div class="card">', unsafe_allow_html=True)
        st.subheader("📝 Extracted Text")

        edited_text = st.text_area(
            "",
            value=st.session_state.extracted_text,
            height=350
        )

        if edited_text:
            st.markdown("### 💾 Export Options")
            d1, d2, d3 = st.columns(3)

            with d1:
                st.download_button(
                    "TXT",
                    edited_text,
                    "extracted_text.txt",
                    "text/plain",
                    use_container_width=True
                )

            with d2:
                word_data = create_word_file(edited_text)
                st.download_button(
                    "Word",
                    word_data,
                    "extracted_text.docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

            with d3:
                pdf_data = create_pdf_file(edited_text)
                st.download_button(
                    "PDF",
                    pdf_data,
                    "application/pdf",
                    use_container_width=True
                )

        st.markdown('</div>', unsafe_allow_html=True)